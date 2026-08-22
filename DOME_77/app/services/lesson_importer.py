from __future__ import annotations
import json,re,shutil,subprocess,tempfile,base64,uuid
from pathlib import Path
from typing import Any
import fitz
import httpx
from docx import Document
from app.core.config import settings
from app.services.authored_content import persistent_lessons_root, backup_lesson_version

SAFE=re.compile(r'[^a-zA-Z0-9_-]+')
URL_RE=re.compile(r'https?://[^\s<>]+')
SLIDE_REF_RE=re.compile(r'(?i)(?:слайд\s*[№#:]?\s*|slide\s*[№#:]?\s*|^\s*)(\d{1,3})(?:\s*(?:[-–—:.]|слайд))?')

def safe_id(v:str)->str:
    x=SAFE.sub('_',v.strip()).strip('_').lower()
    if not x: raise ValueError('lesson_id required')
    return x[:80]

def extract_instruction(path:Path)->str:
    ext=path.suffix.lower()
    if ext=='.docx':
        doc=Document(path)
        paras=[p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Include tables as teacher instructions can be tabular.
        for table in doc.tables:
            for row in table.rows:
                vals=[c.text.strip() for c in row.cells if c.text.strip()]
                if vals: paras.append(' | '.join(vals))
        return '\n'.join(paras)
    if ext=='.pdf':
        doc=fitz.open(path); return '\n'.join(page.get_text('text').strip() for page in doc if page.get_text('text').strip())
    if ext in {'.txt','.md'}: return path.read_text('utf-8',errors='ignore')
    return ''

def to_pdf(path:Path,work:Path)->Path:
    if path.suffix.lower()=='.pdf': return path
    if path.suffix.lower() in {'.pptx','.ppt'}:
        subprocess.run(['libreoffice','--headless','--convert-to','pdf','--outdir',str(work),str(path)],check=True,stdout=subprocess.PIPE,stderr=subprocess.STDOUT,timeout=180)
        out=work/(path.stem+'.pdf')
        if not out.exists(): raise RuntimeError('PPTX conversion failed')
        return out
    raise ValueError('Lesson/homework must be PDF or PPTX')

def render(pdf:Path,out:Path,prefix:str,max_width:int=1100)->list[str]:
    out.mkdir(parents=True,exist_ok=True); doc=fitz.open(pdf); names=[]
    for i,page in enumerate(doc,1):
        scale=max_width/page.rect.width; pix=page.get_pixmap(matrix=fitz.Matrix(scale,scale),alpha=False); name=f'{prefix}_{i:03d}.jpg'; pix.pil_save(str(out/name),format='JPEG',quality=70,optimize=True); names.append(name)
    return names

def instruction_map(text:str)->dict[int,str]:
    """Broad deterministic parser. AI mapping is layered on top when available."""
    lines=[x.rstrip() for x in text.splitlines()]
    hits=[]
    for idx,line in enumerate(lines):
        m=SLIDE_REF_RE.search(line)
        if m and ('слайд' in line.lower() or 'slide' in line.lower() or re.match(r'^\s*\d{1,3}\s*[.)-]',line)):
            hits.append((idx,int(m.group(1))))
    out={}
    for pos,(idx,n) in enumerate(hits):
        end=hits[pos+1][0] if pos+1<len(hits) else len(lines)
        note=' '.join(x.strip() for x in lines[idx:end] if x.strip())
        if note: out[n]=re.sub(r'\s+',' ',note)[:1800]
    return out

def _response_text(payload:dict)->str:
    for item in payload.get('output') or []:
        for c in item.get('content') or []:
            if c.get('type')=='output_text' and c.get('text'): return str(c['text']).strip()
    return ''

def _clean_json(raw:str):
    raw=raw.strip()
    if raw.startswith('```'):
        raw=raw.strip('`').strip()
        if raw.lower().startswith('json'): raw=raw[4:].strip()
    return json.loads(raw)

def instruction_map_ai(text:str, page_count:int)->dict[int,str]:
    if not settings.openai_api_key or not text.strip(): return {}
    headers={'Authorization':f'Bearer {settings.openai_api_key}','Content-Type':'application/json'}
    prompt=(
        'Map this Russian teacher instruction to lesson slide numbers. Return ONLY a JSON object where keys are slide numbers and values are concise complete teacher notes. '
        f'Valid slide numbers are 1..{page_count}. Understand variants such as "3 слайд", "Слайд 3", "3.", tables, and prose references. Do not invent notes. Instruction:\n'+text[:50000]
    )
    try:
        r=httpx.post('https://api.openai.com/v1/responses',headers=headers,json={'model':settings.openai_text_model,'input':prompt},timeout=90); r.raise_for_status()
        obj=_clean_json(_response_text(r.json()))
        return {int(k):str(v)[:1800] for k,v in obj.items() if str(k).isdigit() and 1<=int(k)<=page_count and str(v).strip()}
    except Exception:
        return {}

def infer_type(note:str)->str:
    n=note.lower()
    if 'по рол' in n or 'диалог' in n: return 'read_roles'
    if 'эхо' in n and 'чита' in n: return 'echo_reading'
    if 'чита' in n: return 'read_aloud'
    if ('видео' in n or 'мультф' in n) and any(x in n for x in ['пауза','останов','вопрос после','pause']): return 'video_pause_question'
    if 'мультф' in n or 'видео' in n: return 'video'
    if any(x in n for x in ['обвед','пиши','напиши','нарис','лини']): return 'trace'
    if any(x in n for x in ['перетащ','расстав','собери']): return 'drag_drop'
    if any(x in n for x in ['виктор','выбери','нажми']): return 'choice'
    if '?' in note or any(x in n for x in ['как дума','почему','что произошло','кого ','кто ','куда ','какой ']): return 'comprehension'
    return 'passive'

def analyze_pages_with_ai(image_paths:list[Path], notes:dict[int,str], course_id:str, instruction_text:str="")->dict[int,dict]:
    if not settings.openai_api_key: return {}
    allowed=['passive','video_pause_question','interactive_scene','real_world_find','photo_task','voice_answer','choice','drag_drop','memory','letter_path','trace','tap_sound','match_visible','tap_select','multi_select','matching','sorting','sequence','word_builder','syllable_builder','sentence_builder','fill_gap','odd_one_out','sound_position','syllable_split','find_in_text','connect_lines','handwriting_screen','draw','coloring','maze','dictation','listen_choose','video','read_aloud','read_roles','echo_reading','shared_reading','comprehension','retell','continue_story','mood_choice','physical_action']
    result={}; headers={'Authorization':f'Bearer {settings.openai_api_key}','Content-Type':'application/json'}
    for start in range(0,len(image_paths),4):
        batch=image_paths[start:start+4]
        content=[{'type':'input_text','text':(
            'You are authoring a DOME Russian-language lesson. Analyze each attached slide independently. Course='+course_id+'. Return ONLY JSON array, one object per image in order. '
            'Choose activity_type from: '+','.join(allowed)+'. Respect teacher_note. For interactive slides return concrete executable config. '
            'For drag_drop use items and targets with equal positions representing correct mapping. For matching/match_visible use pairs. For choice use options and correct_indices. '
            'For tap_sound/interactive_scene return hotspots [{label,x,y,w,h}] normalized 0..1. For connect_lines return left_points/right_points normalized 0..1 when detectable. '
            'For maze include start/end normalized points if detectable. For dictation include dictation_text. For video_pause_question include pause_at_seconds and question when known. '
            'For reading use read_aloud/read_roles only when the child is expected to read. For read_aloud/read_roles include exact reading_text. '
            'For read_roles config must prefer ordered turns: role_turns:[{role,text,speaker:"child"|"bot"}], plus child_role when known. Never create Proof of Learning tasks. '
            'The full teacher instruction may be prose and may not name slide numbers. Use its meaning together with the actual slide images to align activities; do not require wording like "5 слайд". '
            'Shape: [{page,activity_type,prompt,audio_text,reading_text,config,confidence}]. Teacher notes already mapped when possible: '+json.dumps({start+i+1:notes.get(start+i+1,'') for i in range(len(batch))},ensure_ascii=False)+
            '. Full teacher instruction: '+str(instruction_text or '')[:30000]
        )}]
        for img in batch:
            b64=base64.b64encode(img.read_bytes()).decode('ascii'); content.append({'type':'input_image','image_url':'data:image/jpeg;base64,'+b64})
        try:
            r=httpx.post('https://api.openai.com/v1/responses',headers=headers,json={'model':settings.openai_text_model,'input':[{'role':'user','content':content}]},timeout=90); r.raise_for_status(); arr=_clean_json(_response_text(r.json()))
            for i,obj in enumerate(arr[:len(batch)]): result[start+i+1]=obj
        except Exception: continue
    return result

def analyze_homework_with_ai(image_paths:list[Path], course_id:str)->dict[int,dict]:
    if not settings.openai_api_key: return {}
    allowed=['trace','handwriting_screen','draw','coloring','maze','choice','tap_select','multi_select','drag_drop','matching','match_visible','connect_lines','word_builder','syllable_builder','sentence_builder','fill_gap','odd_one_out','dictation','tap_sound','comprehension','voice_answer','read_aloud']
    result={}; headers={'Authorization':f'Bearer {settings.openai_api_key}','Content-Type':'application/json'}
    for start in range(0,len(image_paths),4):
        batch=image_paths[start:start+4]
        content=[{'type':'input_text','text':('Analyze DOME Russian-language HOMEWORK pages. Return ONLY JSON array, one object per image. Choose activity_type from: '+','.join(allowed)+'. Prefer trace/handwriting when the child writes on the page. For choices/matching/drag-drop/connect-lines return executable config. Shape: [{page,activity_type,prompt,audio_text,config,confidence}].')}]
        for img in batch:
            b64=base64.b64encode(img.read_bytes()).decode('ascii'); content.append({'type':'input_image','image_url':'data:image/jpeg;base64,'+b64})
        try:
            r=httpx.post('https://api.openai.com/v1/responses',headers=headers,json={'model':settings.openai_text_model,'input':[{'role':'user','content':content}]},timeout=90); r.raise_for_status(); arr=_clean_json(_response_text(r.json()))
            for i,obj in enumerate(arr[:len(batch)]): result[start+i+1]=obj
        except Exception: continue
    return result

def parse_extra_links(texts:list[str]|None)->dict[int,list[str]]:
    out:dict[int,list[str]]={}
    for text in texts or []:
        urls=URL_RE.findall(text)
        if not urls: continue
        m=SLIDE_REF_RE.search(text)
        n=int(m.group(1)) if m else 0
        out.setdefault(n,[]).extend(urls)
    return out

def _slide_number_for_extra(path: Path, note: str = "") -> int:
    """Best-effort mapping for uploaded media. Caption wins; filename is fallback."""
    for text in (note or "", path.stem.replace("_", " ").replace("-", " ")):
        m=SLIDE_REF_RE.search(text)
        if m:
            try: return int(m.group(1))
            except Exception: pass
        m2=re.search(r'(?i)(?:slide|слайд)[ _-]?(\d{1,3})',text)
        if m2: return int(m2.group(1))
    return 0


def import_package(*,lesson_id:str,course_id:str,title:str,lesson_file:Path,instruction_file:Path|None=None,homework_file:Path|None=None,extra_files:list[Path]|None=None,extra_texts:list[str]|None=None,extra_file_notes:list[str]|None=None,target_language:str='ru',order:int=999) -> dict[str,Any]:
    """Build an unpublished content_v1 lesson, then atomically replace the draft/live folder.

    Re-import never deletes the working lesson first. Existing content is versioned and kept
    for rollback. Extra video files can be bound to a slide by Telegram caption or filename,
    e.g. ``слайд 11`` / ``slide_11.mp4``.
    """
    lid=safe_id(lesson_id)
    base=persistent_lessons_root()
    root=base/lid
    stage=base/f'.{lid}.import-{uuid.uuid4().hex[:10]}'
    old_tmp=base/f'.{lid}.old-{uuid.uuid4().hex[:10]}'
    (stage/'source_materials').mkdir(parents=True,exist_ok=False)
    work=Path(tempfile.mkdtemp(prefix='dome_import_'))
    try:
        lesson_copy=stage/'source_materials'/lesson_file.name; shutil.copy2(lesson_file,lesson_copy)
        pdf=to_pdf(lesson_copy,work); images=render(pdf,stage/'images','slide')
        instruction=''; notes={}
        if instruction_file:
            ic=stage/'source_materials'/instruction_file.name; shutil.copy2(instruction_file,ic)
            instruction=extract_instruction(ic)
            (stage/'source_materials'/'instruction.txt').write_text(instruction,'utf-8')
            notes=instruction_map(instruction)
            # Semantic AI mapping supplements deterministic mapping, never erases it.
            notes.update(instruction_map_ai(instruction,len(images)))

        extras_dir=stage/'source_materials'/'extras'; extras_dir.mkdir(exist_ok=True)
        copied_extras=[]; file_map:dict[int,list[str]]={}
        extra_notes=list(extra_file_notes or [])
        for idx,src in enumerate(extra_files or []):
            # avoid overwriting two files that happen to share a name
            name=src.name; dst=extras_dir/name; n=1
            while dst.exists():
                dst=extras_dir/f'{src.stem}_{n}{src.suffix}'; n+=1
            shutil.copy2(src,dst)
            rel=str(dst.relative_to(stage)); copied_extras.append(rel)
            note=extra_notes[idx] if idx<len(extra_notes) else ''
            slide_no=_slide_number_for_extra(src,note)
            file_map.setdefault(slide_no,[]).append(rel)

        link_map=parse_extra_links(extra_texts)
        (stage/'source_materials'/'extra_links.json').write_text(json.dumps({'by_slide':link_map,'raw':extra_texts or []},ensure_ascii=False,indent=2),'utf-8')
        (stage/'source_materials'/'extra_files.json').write_text(json.dumps({'by_slide':file_map,'notes':extra_notes,'files':copied_extras},ensure_ascii=False,indent=2),'utf-8')

        ai_map=analyze_pages_with_ai([stage/'images'/name for name in images],notes,course_id,instruction)
        slides=[]
        video_ext={'.mp4','.webm','.mov','.m4v','.avi','.mkv'}
        for i,name in enumerate(images,1):
            note=notes.get(i,''); ai=ai_map.get(i) or {}
            typ=str(ai.get('activity_type') or infer_type(note)); prompt=str(ai.get('prompt') or note or 'Посмотри на страницу и выполни задание.')
            cfg=ai.get('config') or {}
            expects_answer=typ not in {'passive','video','physical_action'}
            mastery_required=typ in {'drag_drop','trace','handwriting_screen','draw','drawing','dictation','connect_lines','tap_sound','tap_select','multi_select','listen_choose','odd_one_out','find_in_text','match_visible','matching','sorting','sequence','word_builder','syllable_builder','sentence_builder','fill_gap','sound_position','syllable_split','interactive_scene','voice_answer','roleplay','speak','repeat','dialogue','read_aloud','read_roles','echo_reading','shared_reading','comprehension','retell','continue_story'}
            slide={'slide_id':f'slide_{i:03d}','order':i,'type':typ,'title':f'{title} · {i}','prompt':prompt,'audio_text':str(ai.get('audio_text') or prompt),'image_file':f'images/{name}','can_skip':not (expects_answer or mastery_required),'mastery_required':mastery_required,'expects_answer':expects_answer}
            if ai.get('reading_text'): slide['reading_text']=str(ai.get('reading_text'))
            if isinstance(cfg,dict): slide.update(cfg)
            links=link_map.get(i) or []
            if links:
                slide['video_url']=links[0]
                slide['attachments']=list(dict.fromkeys((slide.get('attachments') or [])+links))
                if typ=='passive': slide['type']='video'
            attachments=file_map.get(i) or []
            if attachments:
                slide['attachments']=list(dict.fromkeys((slide.get('attachments') or [])+attachments))
                video_rel=next((rel for rel in attachments if Path(rel).suffix.lower() in video_ext),None)
                if video_rel:
                    slide['video_file']=video_rel
                    if slide['type']=='passive': slide['type']='video'
            # Any unassigned extras remain visible at lesson level; do not silently lose them.
            slide['authoring_confidence']=float(ai.get('confidence') or (0.55 if note else 0.35))
            slides.append(slide)

        lesson={'schema_version':'1.3','engine':'content_v1','lesson_id':lid,'course_id':course_id,'title':title,'order':int(order),'active':False,'target_language':target_language,'target_duration_minutes':35,'duration_range_minutes':[30,40],'max_completed_runs':2,'expires_after_months':10,'adaptive_second_run':True,'make_cartoon':course_id=='conversation','cartoon_on_first_run_only':True,'slides':slides,'extra_files':copied_extras,'extra_links':[u for values in link_map.values() for u in values],'unassigned_extra_files':file_map.get(0,[]),'unassigned_extra_links':link_map.get(0,[]),'import_status':'REVIEW_REQUIRED'}
        (stage/'lesson.json').write_text(json.dumps(lesson,ensure_ascii=False,indent=2),'utf-8')

        if homework_file:
            hc=stage/'source_materials'/homework_file.name; shutil.copy2(homework_file,hc)
            hpdf=to_pdf(hc,work); himages=render(hpdf,stage/'homework_images','hw')
            hw_ai=analyze_homework_with_ai([stage/'homework_images'/name for name in himages],course_id)
            hslides=[]
            for i,name in enumerate(himages,1):
                ai=hw_ai.get(i) or {}; typ=str(ai.get('activity_type') or 'trace'); cfg=ai.get('config') or {}; prompt=str(ai.get('prompt') or 'Выполни задание прямо на странице. Можно писать, обводить и рисовать пальцем или стилусом.')
                hs={'slide_id':f'hw_{i:02d}','order':i,'type':typ,'title':f'Домашнее задание · {i}','prompt':prompt,'audio_text':str(ai.get('audio_text') or prompt),'image_file':f'homework_images/{name}','can_skip':False,'mastery_required':True,'expects_answer':True,'authoring_confidence':float(ai.get('confidence') or 0.35)}
                if isinstance(cfg,dict): hs.update(cfg)
                hslides.append(hs)
            hw={'schema_version':'1.3','title':f'Домашнее задание · {title}','optional':True,'duration_minutes':10,'send_to_bot':True,'send_to_parent_email':True,'source_file':f'source_materials/{homework_file.name}','slides':hslides}
            (stage/'homework.json').write_text(json.dumps(hw,ensure_ascii=False,indent=2),'utf-8')

        # Only now touch the existing lesson. Preserve its version history in the new folder.
        if root.exists() and (root/'lesson.json').exists():
            backup_lesson_version(lid,'before_reimport')
            if (root/'_versions').exists(): shutil.copytree(root/'_versions',stage/'_versions',dirs_exist_ok=True)
        if root.exists(): root.rename(old_tmp)
        try:
            stage.rename(root)
        except Exception:
            if old_tmp.exists() and not root.exists(): old_tmp.rename(root)
            raise
        shutil.rmtree(old_tmp,ignore_errors=True)
        return lesson
    finally:
        shutil.rmtree(work,ignore_errors=True)
        if stage.exists(): shutil.rmtree(stage,ignore_errors=True)
        if old_tmp.exists() and root.exists(): shutil.rmtree(old_tmp,ignore_errors=True)

